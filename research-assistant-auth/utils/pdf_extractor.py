import fitz
import pdfplumber
import docx
import os
from text_cleaner import clean_text


def extract_pdf_text(file_path: str):

    text_data = []

    try:

        with pdfplumber.open(file_path) as pdf:

            for page_num, page in enumerate(

                pdf.pages,

                start=1

            ):

                text = page.extract_text()

                # -----------------------
                # Quality Validation
                # -----------------------

                if not text:

                    text = ""

                    heading = ""

                    scanned = True

                else:

                    heading = text.split(

                        "\n"

                    )[0].strip()

                    scanned = False

                text_data.append(

                    {

                        "file_id": "",

                        "heading": heading,

                        "page": page_num,

                        "text": text,

                        "blank_page":

                        len(text.strip()) == 0,

                        "scanned_page":

                        scanned

                    }

                )

    except:

        pdf = fitz.open(file_path)

        for page_num in range(

            len(pdf)

        ):

            page = pdf[page_num]

            page_text = page.get_text()

            if not page_text:

                page_text = ""

                heading = ""

                scanned = True

            else:

                heading = page_text.split(

                    "\n"

                )[0].strip()

                scanned = False

            text_data.append(

                {

                    "file_id": "",

                    "heading": heading,

                    "page": page_num + 1,

                    "text": page_text,

                    "blank_page":

                    len(

                        page_text.strip()

                    ) == 0,

                    "scanned_page":

                    scanned

                }

            )

    return text_data


def extract_docx_text(file_path: str):

    doc = docx.Document(file_path)

    full_text = "\n".join(

        [

            para.text

            for para in doc.paragraphs

            if para.text.strip() != ""

        ]

    )

    return [

        {

            "file_id": "",

            "heading": full_text.split("\n")[0].strip() if full_text else "",

            "page": 1,

            "text": full_text

        }

    ]


def extract_txt_text(file_path: str):

    with open(

        file_path,

        "r",

        encoding="utf-8"

    ) as f:

        text = f.read()

    return [

        {

            "file_id": "",

            "heading": text.split("\n")[0].strip() if text else "",

            "page": 1,

            "text": text

        }

    ]


def extract_document(

    file_path: str,

    file_type: str

):

    file_type = file_type.lower()

    if file_type == "pdf":

        extracted_data = extract_pdf_text(

            file_path

        )

    elif file_type == "docx":

        extracted_data = extract_docx_text(

            file_path

        )

    elif file_type == "txt":

        extracted_data = extract_txt_text(

            file_path

        )

    else:

        raise Exception(

            "Unsupported file type"

        )


    for item in extracted_data:

        item["text"] = clean_text(

            item["text"]

        )

    return extracted_data
    

def extract_metadata(

    file_path: str,

    file_type: str

):

    metadata = {

        "title": "",

        "authors": "",

        "year": "",

        "page_count": 0,

        "file_size": os.path.getsize(file_path),

        "keywords": "",

        "doi": ""

    }

    if file_type.lower() == "pdf":

        pdf = fitz.open(file_path)

        info = pdf.metadata

        metadata["title"] = info.get("title", "")

        metadata["authors"] = info.get("author", "")

        metadata["page_count"] = len(pdf)

        metadata["keywords"] = info.get("keywords", "")

        metadata["doi"] = ""

        metadata["year"] = info.get("creationDate", "")

        pdf.close()

    elif file_type.lower() == "docx":

        doc = docx.Document(file_path)

        props = doc.core_properties

        metadata["title"] = props.title

        metadata["authors"] = props.author

        metadata["year"] = str(props.created) if props.created else ""

        metadata["page_count"] = 1

    elif file_type.lower() == "txt":

        metadata["title"] = os.path.basename(file_path)

        metadata["page_count"] = 1

    return metadata